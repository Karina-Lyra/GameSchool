import os
import json
from tkinter import filedialog, messagebox
import docx
from PyPDF2 import PdfReader


# Função para extrair perguntas de um arquivo de texto
def importar_txt(caminho_arquivo):
    perguntas = []
    with open(caminho_arquivo, 'r', encoding='utf-8') as f:
        conteudo = f.read().splitlines()
        for i in range(0, len(conteudo), 2):  # Considera pergunta na linha ímpar e resposta na linha seguinte
            if i + 1 < len(conteudo):
                pergunta = conteudo[i]
                resposta = conteudo[i + 1]
                perguntas.append({'pergunta': pergunta, 'resposta': resposta})
    return perguntas


# Função para extrair perguntas de um arquivo PDF
def importar_pdf(caminho_arquivo):
    perguntas = []
    with open(caminho_arquivo, 'rb') as f:
        reader = PdfReader(f)
        for page in reader.pages:
            texto = page.extract_text()
            linhas = texto.split('\n')
            for i in range(0, len(linhas), 2):  # Considera pergunta na linha ímpar e resposta na linha seguinte
                if i + 1 < len(linhas):
                    pergunta = linhas[i]
                    resposta = linhas[i + 1]
                    perguntas.append({'pergunta': pergunta, 'resposta': resposta})
    return perguntas


# Função para extrair perguntas de um arquivo DOCX
def importar_docx(caminho_arquivo):
    perguntas = []
    doc = docx.Document(caminho_arquivo)
    for i in range(0, len(doc.paragraphs), 2):  # Considera pergunta na linha ímpar e resposta na linha seguinte
        if i + 1 < len(doc.paragraphs):
            pergunta = doc.paragraphs[i].text
            resposta = doc.paragraphs[i + 1].text
            perguntas.append({'pergunta': pergunta, 'resposta': resposta})
    return perguntas


# Função para importar um arquivo de pergunta
def importar_arquivo():
    caminho_arquivo = filedialog.askopenfilename(title="Escolha o arquivo de perguntas", filetypes=(
    ("Text Files", "*.txt"), ("PDF Files", "*.pdf"), ("Word Files", "*.docx")))
    if not caminho_arquivo:
        return

    if caminho_arquivo.endswith('.txt'):
        perguntas = importar_txt(caminho_arquivo)
    elif caminho_arquivo.endswith('.pdf'):
        perguntas = importar_pdf(caminho_arquivo)
    elif caminho_arquivo.endswith('.docx'):
        perguntas = importar_docx(caminho_arquivo)
    else:
        messagebox.showerror("Erro", "Tipo de arquivo não suportado.")
        return

    # Salvar as perguntas importadas em um arquivo JSON
    nome_arquivo = os.path.splitext(os.path.basename(caminho_arquivo))[0] + '.json'
    caminho_json = os.path.join('banco_perguntas', nome_arquivo)

    with open(caminho_json, 'w', encoding='utf-8') as f:
        json.dump(perguntas, f, ensure_ascii=False, indent=4)

    messagebox.showinfo("Importação concluída", f"As perguntas foram importadas com sucesso para {nome_arquivo}!")

